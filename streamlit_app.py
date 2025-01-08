import streamlit as st
import faiss
import os
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEndpoint
from langchain.prompts import PromptTemplate
import time

from hf_key import hf_token_key  # Set the Hugging Face Hub API token as an environment variable
os.environ['HUGGINGFACEHUB_API_TOKEN'] = hf_token_key

def process_input(input_type, input_data):
    loader = None
    if input_type == "Web Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for PDF")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data  # Input is already a text string
        else:
            raise ValueError("Expected a string for 'Text' input type.")
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for DOCX")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    else:
        raise ValueError("Unsupported input type")

    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == "Web Link":
        texts = text_splitter.split_documents(documents)
        texts = [ str(doc.page_content) for doc in texts ]  # Access page_content from each Document 
    else:
        texts = text_splitter.split_text(documents)

    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}

    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )
    # Create FAISS index
    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)
    # Create FAISS vector store with the embedding function
    vector_store = FAISS(
        embedding_function=hf_embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)  # Add documents to the vector store
    return vector_store



def answer_question(vectorstore, query):
    """Answers a question based on the provided vectorstore."""
    qa_template = """Use the given context to answer the question.
    If you don't know the answer, just say that you don't know, don't try to make up an answer.
    Use six sentences maximum and keep the answer concise. 

    Context: {context}

    Question: {question}
    Answer:
    """
    # Define the prompt template for Q&A
    qa_prompt_template = PromptTemplate.from_template(qa_template)
    llm = HuggingFaceEndpoint(repo_id= 'meta-llama/Meta-Llama-3-8B-Instruct', 
                              token = hf_token_key, temperature= 0.2)
    qa = RetrievalQA.from_chain_type(llm=llm,
                                      retriever=vectorstore.as_retriever(),
                                      chain_type = 'stuff',
                                      chain_type_kwargs={"prompt": qa_prompt_template}
                                      )
    
    #return qa
    result = qa({"query": query})
    # Assume vectorstore is already initialized
    

    return  result["result"]


def stream_answer(answer):
    """Simulates streaming text output."""
    for chunk in answer.split(" "):  # Split the answer into chunks (sentences)
        yield chunk + " "  
        time.sleep(0.1)



def main():
    st.title("RAG Llama 3 🦙 Q&A ")
    input_type = st.selectbox("Input Type", 
                              ("Web Link", "PDF", "Text", "DOCX"),
                              index = None,
                              placeholder= "Choose input type for Q/A"
                              )
    
    if input_type == "Web Link":
        input_data = st.text_input("Enter URL link",
                                   placeholder = "https://www.example.com")
    elif input_type == "Text":
        input_data = st.text_input("Enter text", placeholder='Paste text here')
    elif input_type == 'PDF':
        input_data = st.file_uploader("Upload a .PDF file", type=["pdf"])
    elif input_type == 'DOCX':
        input_data = st.file_uploader("Upload a .DOC/.DOCX file", type=[ 'docx', 'doc'])
    if st.button("Proceed"):
        # st.write(process_input(input_type, input_data))
        vectorstore = process_input(input_type, input_data)
        st.session_state["vectorstore"] = vectorstore
    if "vectorstore" in st.session_state:
        query = st.text_input("Ask your question")
        if st.button("Generate Answer"):
            answer = answer_question(st.session_state["vectorstore"], query)
            st.subheader("Answer:")
            st.write(stream_answer(answer))
            #st.text(answer) 
            st.divider()
            st.text("You can ask a new question again from the same RAG input")
            #st.write(answer)

if __name__ == "__main__":
    main()